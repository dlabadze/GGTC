from odoo import models, fields, api
from odoo.exceptions import UserError
import base64
import io
import logging

_logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    _logger.warning("python-docx library not available. Word document signing will not work.")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    _logger.warning("PIL/Pillow library not available. Image conversion may not work.")

try:
    import cairosvg
    CAIROSVG_AVAILABLE = True
except ImportError:
    CAIROSVG_AVAILABLE = False
    _logger.warning("cairosvg library not available. SVG signature conversion may not work.")

class ApprovalRequest(models.Model):
    _inherit = 'approval.request'
    signed_document = fields.Boolean(string="ხელმოწერილი", default=False)
    request_from = fields.Many2one("approval.request", string='Request From')
    related_request_ids = fields.One2many("approval.request", "request_from", string='Related Requests')
    root_request_id = fields.Many2one(
        "approval.request",
        string='Root Request',
        compute='_compute_root_request',
        store=False,
    )

    @api.depends('request_from', 'request_from.root_request_id')
    def _compute_root_request(self):
        for request in self:
            root = request
            visited = self.env['approval.request']
            while root.request_from and root.request_from not in visited:
                visited |= root
                root = root.request_from
            request.root_request_id = root if root != request else False

    def action_create_new_request(self):
        return {
            'type': 'ir.actions.act_window',
            'name': 'Create New Request',
            'res_model': 'create.new.request.wizard',
            'view_mode': 'form',
            'target': 'new',
        }

    def action_view_related_request(self):
        return {
            'type': 'ir.actions.act_window',
            'name': 'Approval Request',
            'res_model': 'approval.request',
            'res_id': self.id,
            'view_mode': 'form',
            'target': 'current',
        }

    def action_open_related_document(self):
        self.ensure_one()

        if self.request_from:
            view_id = self.env['ir.ui.view'].search([
                ('model', '=', 'approval.request'),
                ('type', '=', 'form')
            ], limit=1)

            return {
                'type': 'ir.actions.act_window',
                'res_model': 'approval.request',
                'res_id': self.request_from.id,
                'view_mode': 'form',
                'view_id': view_id.id if view_id else False,
                'target': 'current',
            }

    def action_open_root_request(self):
        self.ensure_one()
        root_request = self.root_request_id

        if not root_request:
            return self.action_view_related_request()

        view_id = self.env['ir.ui.view'].search([
            ('model', '=', 'approval.request'),
            ('type', '=', 'form')
        ], limit=1)

        return {
            'type': 'ir.actions.act_window',
            'res_model': 'approval.request',
            'res_id': root_request.id,
            'view_mode': 'form',
            'view_id': view_id.id if view_id else False,
            'target': 'current',
        }

    def action_sign_documents(self):
        """Sign Word documents by adding user signature to the last page right side"""
        self.ensure_one()
        
        if not DOCX_AVAILABLE:
            raise UserError("python-docx library is not installed. Please install it to use document signing.")
        
        # Get user signature
        user = self.env.user
        if not user.sign_signature:
            raise UserError("You don't have a signature configured. Please set your signature first.")
        
        # Get all attachments
        attachments = self.env['ir.attachment'].search([
            ('res_model', '=', 'approval.request'),
            ('res_id', '=', self.id)
        ])
        
        if not attachments:
            raise UserError("No attachments found for this approval request.")
        
        # Filter Word documents
        word_attachments = attachments.filtered(
            lambda att: att.mimetype == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' 
            or att.name and att.name.lower().endswith('.docx')
        )
        
        if not word_attachments:
            raise UserError("No Word document attachments found (.docx files).")
        
        # Decode signature image
        try:
            signature_value = user.sign_signature
            # Handle bytes or string format
            if isinstance(signature_value, bytes):
                signature_data = signature_value
            else:
                # Handle data URI format (data:image/png;base64,...)
                if signature_value.startswith('data:'):
                    signature_data = base64.b64decode(signature_value.split(',')[1])
                else:
                    signature_data = base64.b64decode(signature_value)
            
            # Validate signature data
            if not signature_data:
                raise UserError("Signature image is empty.")
            
            # Check if it's a valid image (basic check - PNG, JPEG, etc.)
            if len(signature_data) < 10:
                raise UserError("Signature image appears to be corrupted or invalid.")
            
            # Detect and convert the image format
            # Check if it's SVG (common in Odoo signatures)
            signature_str = signature_data.decode('utf-8', errors='ignore') if isinstance(signature_data, bytes) else str(signature_data)
            is_svg = signature_str.strip().startswith('<svg') or signature_str.strip().startswith('<?xml')
            
            if is_svg and CAIROSVG_AVAILABLE:
                # Convert SVG to PNG
                try:
                    png_data = cairosvg.svg2png(bytestring=signature_data)
                    signature_data = png_data
                    _logger.info("SVG signature converted to PNG format")
                except Exception as svg_error:
                    _logger.error(f"Error converting SVG signature: {str(svg_error)}", exc_info=True)
                    raise UserError(f"Error converting SVG signature to image. Please update your signature. Error: {str(svg_error)}")
            elif is_svg and not CAIROSVG_AVAILABLE:
                raise UserError("Signature is in SVG format but cairosvg library is not installed. Please install it: pip install cairosvg")
            
            # Try to validate and convert the image format using PIL if available
            if PIL_AVAILABLE:
                try:
                    # Check if we can identify the image format
                    img_stream = io.BytesIO(signature_data)
                    img_stream.seek(0)
                    
                    # Check the first bytes to identify format
                    header = signature_data[:16] if len(signature_data) >= 16 else signature_data
                    _logger.debug(f"Image header (hex): {header.hex() if isinstance(header, bytes) else 'N/A'}")
                    
                    # Try to open and identify the image
                    try:
                        img = Image.open(img_stream)
                        # Verify it's a valid image by trying to load it
                        img.verify()
                        # Reset stream after verify (verify() closes the image)
                        img_stream.seek(0)
                        img = Image.open(img_stream)
                    except Exception as open_error:
                        # If we can't open it, check if it might be base64 encoded again
                        try:
                            # Maybe it's double-encoded?
                            if isinstance(signature_data, bytes):
                                decoded = base64.b64decode(signature_data)
                                if decoded != signature_data:
                                    img_stream = io.BytesIO(decoded)
                                    img_stream.seek(0)
                                    img = Image.open(img_stream)
                                    img.verify()
                                    img_stream.seek(0)
                                    img = Image.open(img_stream)
                                    signature_data = decoded
                                else:
                                    raise open_error
                            else:
                                raise open_error
                        except:
                            raise open_error
                    
                    # Convert to RGB PNG format for maximum compatibility
                    if img.mode in ('RGBA', 'LA', 'P'):
                        rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'P':
                            img = img.convert('RGBA')
                        if img.mode in ('RGBA', 'LA'):
                            rgb_img.paste(img, mask=img.split()[-1])
                        else:
                            rgb_img.paste(img)
                        img = rgb_img
                    elif img.mode != 'RGB':
                        img = img.convert('RGB')
                    # Save as PNG in memory
                    png_stream = io.BytesIO()
                    img.save(png_stream, format='PNG')
                    png_stream.seek(0)
                    signature_data = png_stream.read()
                    _logger.info("Signature image successfully converted to PNG format")
                except Exception as img_error:
                    error_msg = str(img_error)
                    _logger.error(f"Error processing signature image: {error_msg}", exc_info=True)
                    # Check the first few bytes to identify format
                    img_header = signature_data[:20] if len(signature_data) >= 20 else signature_data
                    header_info = img_header.hex() if isinstance(img_header, bytes) else str(img_header)[:50]
                    _logger.error(f"Signature data header (hex): {header_info}")
                    raise UserError(f"Signature image format is not recognized or corrupted. Please update your signature in user settings. Error: {error_msg}")
                
        except UserError:
            raise
        except Exception as e:
            raise UserError(f"Error decoding signature: {str(e)}")
        
        # Process each Word document
        for attachment in word_attachments:
            try:
                # Read the document
                doc_data = base64.b64decode(attachment.datas)
                if not doc_data:
                    raise UserError(f"Document '{attachment.name}' is empty or corrupted.")
                
                doc_stream = io.BytesIO(doc_data)
                doc = Document(doc_stream)
                
                # Search for "[ხელმოწერა]" text and replace with signature image
                signature_placeholder = "[ხელმოწერა]"
                signature_found = False
                
                # Iterate through all paragraphs to find the placeholder text
                for paragraph in doc.paragraphs:
                    paragraph_text = paragraph.text
                    if signature_placeholder in paragraph_text:
                        signature_found = True
                        # Split the text to preserve text before and after the placeholder
                        text_parts = paragraph_text.split(signature_placeholder)
                        # Clear the paragraph and rebuild it with the signature image
                        paragraph.clear()
                        if len(text_parts) > 0:
                            # Add text before placeholder if exists
                            if text_parts[0].strip():
                                paragraph.add_run(text_parts[0])
                        
                        # Add the signature image inline
                        signature_stream = io.BytesIO(signature_data)
                        run = paragraph.add_run()
                        
                        try:
                            signature_stream.seek(0)
                            # Use height instead of width to keep it inline with text (smaller size)
                            # Height of 0.6 inches should keep it on the same line
                            run.add_picture(signature_stream, height=Inches(0.6))
                        except Exception as img_error:
                            error_type = type(img_error).__name__
                            error_details = str(img_error) or error_type
                            _logger.error(f"Error adding signature image to {attachment.name}: {error_type} - {error_details}", exc_info=True)
                            
                            # Provide more specific error messages
                            if 'image' in error_details.lower() or 'format' in error_details.lower() or 'UnrecognizedImageError' in error_type:
                                raise UserError(f"Signature image format not supported for '{attachment.name}'. Please update your signature in user settings. Error: {error_details}")
                            elif 'BadZipFile' in error_type or 'zip' in error_details.lower():
                                raise UserError(f"Invalid image file format for signature. Error: {error_details}")
                            else:
                                raise UserError(f"Error adding signature image to '{attachment.name}': {error_type} - {error_details}")
                        
                        # Add text after placeholder if exists
                        if len(text_parts) > 1 and text_parts[1].strip():
                            paragraph.add_run(text_parts[1])
                        
                        break  # Only replace the first occurrence
                
                # Also search in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph_text = paragraph.text
                                if signature_placeholder in paragraph_text:
                                    signature_found = True
                                    # Split the text to preserve text before and after the placeholder
                                    text_parts = paragraph_text.split(signature_placeholder)
                                    # Clear and rebuild paragraph with signature
                                    paragraph.clear()
                                    if len(text_parts) > 0 and text_parts[0].strip():
                                        paragraph.add_run(text_parts[0])
                                    
                                    # Add signature image
                                    signature_stream = io.BytesIO(signature_data)
                                    run = paragraph.add_run()
                                    
                                    try:
                                        signature_stream.seek(0)
                                        run.add_picture(signature_stream, height=Inches(0.6))
                                    except Exception as img_error:
                                        error_type = type(img_error).__name__
                                        error_details = str(img_error) or error_type
                                        _logger.error(f"Error adding signature image in table to {attachment.name}: {error_type} - {error_details}", exc_info=True)
                                        raise UserError(f"Error adding signature image to '{attachment.name}': {error_details}")
                                    
                                    if len(text_parts) > 1 and text_parts[1].strip():
                                        paragraph.add_run(text_parts[1])
                                    
                                    break
                
                if not signature_found:
                    raise UserError(f"Signature placeholder '[ხელმოწერა]' not found in document '{attachment.name}'. Please add '[ხელმოწერა]' where you want the signature to appear.")
                
                # Save the modified document
                output_stream = io.BytesIO()
                doc.save(output_stream)
                output_stream.seek(0)
                
                # Update the attachment
                output_data = output_stream.read()
                if not output_data:
                    raise UserError(f"Failed to generate signed document for '{attachment.name}'.")
                
                attachment.write({
                    'datas': base64.b64encode(output_data).decode('utf-8')
                })
                
            except UserError:
                # Re-raise UserError as-is
                raise
            except Exception as e:
                error_msg = str(e) or type(e).__name__
                error_type = type(e).__name__
                _logger.error(f"Error processing attachment {attachment.name}: {error_type} - {error_msg}", exc_info=True)
                if not error_msg or error_msg == error_type:
                    raise UserError(f"Error processing document '{attachment.name}': {error_type}. Please check if the document is a valid Word document (.docx) and try again.")
                else:
                    raise UserError(f"Error processing document '{attachment.name}': {error_msg}")
        self.write({'signed_document': True})

        return {
            'type': 'ir.actions.client',
            'tag': 'display_notification',
            'params': {
                'title': 'Success',
                'message': f'Successfully signed {len(word_attachments)} document(s).',
                'type': 'success',
                'sticky': False,
                'next': {
                    'type': 'ir.actions.client',
                    'tag': 'reload',
                }
            }
        }

    def action_sign_documents_2(self):
        """Sign Word document from x_studio_file field by adding user signature to the last page right side"""
        self.ensure_one()
        
        if not DOCX_AVAILABLE:
            raise UserError("python-docx library is not installed. Please install it to use document signing.")
        
        # Get user signature
        user = self.env.user
        if not user.sign_signature:
            raise UserError("You don't have a signature configured. Please set your signature first.")
        
        # Check if x_studio_file field has data
        if not self.x_studio_file:
            raise UserError("No file found in x_studio_file field.")
        
        # Decode signature image
        try:
            signature_value = user.sign_signature
            # Handle bytes or string format
            if isinstance(signature_value, bytes):
                signature_data = signature_value
            else:
                # Handle data URI format (data:image/png;base64,...)
                if signature_value.startswith('data:'):
                    signature_data = base64.b64decode(signature_value.split(',')[1])
                else:
                    signature_data = base64.b64decode(signature_value)
            
            # Validate signature data
            if not signature_data:
                raise UserError("Signature image is empty.")
            
            # Check if it's a valid image (basic check - PNG, JPEG, etc.)
            if len(signature_data) < 10:
                raise UserError("Signature image appears to be corrupted or invalid.")
            
            # Detect and convert the image format
            # Check if it's SVG (common in Odoo signatures)
            signature_str = signature_data.decode('utf-8', errors='ignore') if isinstance(signature_data, bytes) else str(signature_data)
            is_svg = signature_str.strip().startswith('<svg') or signature_str.strip().startswith('<?xml')
            
            if is_svg and CAIROSVG_AVAILABLE:
                # Convert SVG to PNG
                try:
                    png_data = cairosvg.svg2png(bytestring=signature_data)
                    signature_data = png_data
                    _logger.info("SVG signature converted to PNG format")
                except Exception as svg_error:
                    _logger.error(f"Error converting SVG signature: {str(svg_error)}", exc_info=True)
                    raise UserError(f"Error converting SVG signature to image. Please update your signature. Error: {str(svg_error)}")
            elif is_svg and not CAIROSVG_AVAILABLE:
                raise UserError("Signature is in SVG format but cairosvg library is not installed. Please install it: pip install cairosvg")
            
            # Try to validate and convert the image format using PIL if available
            if PIL_AVAILABLE:
                try:
                    # Check if we can identify the image format
                    img_stream = io.BytesIO(signature_data)
                    img_stream.seek(0)
                    
                    # Check the first bytes to identify format
                    header = signature_data[:16] if len(signature_data) >= 16 else signature_data
                    _logger.debug(f"Image header (hex): {header.hex() if isinstance(header, bytes) else 'N/A'}")
                    
                    # Try to open and identify the image
                    try:
                        img = Image.open(img_stream)
                        # Verify it's a valid image by trying to load it
                        img.verify()
                        # Reset stream after verify (verify() closes the image)
                        img_stream.seek(0)
                        img = Image.open(img_stream)
                    except Exception as open_error:
                        # If we can't open it, check if it might be base64 encoded again
                        try:
                            # Maybe it's double-encoded?
                            if isinstance(signature_data, bytes):
                                decoded = base64.b64decode(signature_data)
                                if decoded != signature_data:
                                    img_stream = io.BytesIO(decoded)
                                    img_stream.seek(0)
                                    img = Image.open(img_stream)
                                    img.verify()
                                    img_stream.seek(0)
                                    img = Image.open(img_stream)
                                    signature_data = decoded
                                else:
                                    raise open_error
                            else:
                                raise open_error
                        except:
                            raise open_error
                    
                    # Convert to RGB PNG format for maximum compatibility
                    if img.mode in ('RGBA', 'LA', 'P'):
                        rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'P':
                            img = img.convert('RGBA')
                        if img.mode in ('RGBA', 'LA'):
                            rgb_img.paste(img, mask=img.split()[-1])
                        else:
                            rgb_img.paste(img)
                        img = rgb_img
                    elif img.mode != 'RGB':
                        img = img.convert('RGB')
                    # Save as PNG in memory
                    png_stream = io.BytesIO()
                    img.save(png_stream, format='PNG')
                    png_stream.seek(0)
                    signature_data = png_stream.read()
                    _logger.info("Signature image successfully converted to PNG format")
                except Exception as img_error:
                    error_msg = str(img_error)
                    _logger.error(f"Error processing signature image: {error_msg}", exc_info=True)
                    # Check the first few bytes to identify format
                    img_header = signature_data[:20] if len(signature_data) >= 20 else signature_data
                    header_info = img_header.hex() if isinstance(img_header, bytes) else str(img_header)[:50]
                    _logger.error(f"Signature data header (hex): {header_info}")
                    raise UserError(f"Signature image format is not recognized or corrupted. Please update your signature in user settings. Error: {error_msg}")
                
        except UserError:
            raise
        except Exception as e:
            raise UserError(f"Error decoding signature: {str(e)}")
        
        # Process the document from x_studio_file
        try:
            # Read the document from x_studio_file field
            doc_data = base64.b64decode(self.x_studio_file)
            if not doc_data:
                raise UserError("Document in x_studio_file field is empty or corrupted.")
            
            # Check if it's a Word document (basic check)
            # Word documents start with PK (ZIP file signature)
            if not doc_data.startswith(b'PK'):
                raise UserError("The file in x_studio_file field does not appear to be a valid Word document (.docx file).")
            
            doc_stream = io.BytesIO(doc_data)
            doc = Document(doc_stream)
            
            # Search for "[ხელმოწერა]" text and replace with signature image
            signature_placeholder = "[ხელმოწერა]"
            signature_found = False
            
            # Iterate through all paragraphs to find the placeholder text
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text
                if signature_placeholder in paragraph_text:
                    signature_found = True
                    # Split the text to preserve text before and after the placeholder
                    text_parts = paragraph_text.split(signature_placeholder)
                    # Clear the paragraph and rebuild it with the signature image
                    paragraph.clear()
                    if len(text_parts) > 0:
                        # Add text before placeholder if exists
                        if text_parts[0].strip():
                            paragraph.add_run(text_parts[0])
                    
                    # Add the signature image inline
                    signature_stream = io.BytesIO(signature_data)
                    run = paragraph.add_run()
                    
                    try:
                        signature_stream.seek(0)
                        # Use height instead of width to keep it inline with text (smaller size)
                        # Height of 0.6 inches should keep it on the same line
                        run.add_picture(signature_stream, height=Inches(0.6))
                    except Exception as img_error:
                        error_type = type(img_error).__name__
                        error_details = str(img_error) or error_type
                        _logger.error(f"Error adding signature image: {error_type} - {error_details}", exc_info=True)
                        
                        # Provide more specific error messages
                        if 'image' in error_details.lower() or 'format' in error_details.lower() or 'UnrecognizedImageError' in error_type:
                            raise UserError(f"Signature image format not supported. Please update your signature in user settings. Error: {error_details}")
                        elif 'BadZipFile' in error_type or 'zip' in error_details.lower():
                            raise UserError(f"Invalid image file format for signature. Error: {error_details}")
                        else:
                            raise UserError(f"Error adding signature image: {error_type} - {error_details}")
                    
                    # Add text after placeholder if exists
                    if len(text_parts) > 1 and text_parts[1].strip():
                        paragraph.add_run(text_parts[1])
                    
                    break  # Only replace the first occurrence
            
            # Also search in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph_text = paragraph.text
                            if signature_placeholder in paragraph_text:
                                signature_found = True
                                # Split the text to preserve text before and after the placeholder
                                text_parts = paragraph_text.split(signature_placeholder)
                                # Clear and rebuild paragraph with signature
                                paragraph.clear()
                                if len(text_parts) > 0 and text_parts[0].strip():
                                    paragraph.add_run(text_parts[0])
                                
                                # Add signature image
                                signature_stream = io.BytesIO(signature_data)
                                run = paragraph.add_run()
                                
                                try:
                                    signature_stream.seek(0)
                                    run.add_picture(signature_stream, height=Inches(0.6))
                                except Exception as img_error:
                                    error_type = type(img_error).__name__
                                    error_details = str(img_error) or error_type
                                    _logger.error(f"Error adding signature image in table: {error_type} - {error_details}", exc_info=True)
                                    raise UserError(f"Error adding signature image: {error_details}")
                                
                                if len(text_parts) > 1 and text_parts[1].strip():
                                    paragraph.add_run(text_parts[1])
                                
                                break
            
            if not signature_found:
                raise UserError("Signature placeholder '[ხელმოწერა]' not found in document. Please add '[ხელმოწერა]' where you want the signature to appear.")
            
            # Save the modified document
            output_stream = io.BytesIO()
            doc.save(output_stream)
            output_stream.seek(0)
            
            # Update the x_studio_file field
            output_data = output_stream.read()
            if not output_data:
                raise UserError("Failed to generate signed document.")
            
            self.write({
                'x_studio_file': base64.b64encode(output_data).decode('utf-8'),
                'signed_document': True
            })
            for req in self:
                current_user = self.env.user
                # Mark activities as done for the current user
                req.activity_ids.filtered(lambda a: a.user_id == current_user).action_feedback(feedback="ხელმოწერილია")
                action = self.env['ir.actions.server'].browse(1716)
                if action.exists():
                    action.with_context(active_id=req.id, active_model=req._name).run()
            
        except UserError:
            # Re-raise UserError as-is
            raise
        except Exception as e:
            error_msg = str(e) or type(e).__name__
            error_type = type(e).__name__
            _logger.error(f"Error processing document from x_studio_file: {error_type} - {error_msg}", exc_info=True)
            if not error_msg or error_msg == error_type:
                raise UserError(f"Error processing document: {error_type}. Please check if the document is a valid Word document (.docx) and try again.")
            else:
                raise UserError(f"Error processing document: {error_msg}")
        
        return {
            'type': 'ir.actions.client',
            'tag': 'display_notification',
            'params': {
                'title': 'Success',
                'message': 'Successfully signed document.',
                'type': 'success',
                'sticky': False,
                'next': {
                    'type': 'ir.actions.client',
                    'tag': 'reload',
                }
            }
        }