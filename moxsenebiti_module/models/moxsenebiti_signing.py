from odoo import models, fields, _
from odoo.exceptions import UserError
import base64
import io
import logging

_logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    _logger.warning("python-docx library not available. Word document signing will not work.")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    _logger.warning("PIL/Pillow library not available. Image conversion may not work.")

try:
    import cairosvg
    CAIROSVG_AVAILABLE = True
except ImportError:
    CAIROSVG_AVAILABLE = False
    _logger.warning("cairosvg library not available. SVG signature conversion may not work.")


class Moxsenebiti(models.Model):
    _inherit = "moxsenebiti"

    def _sign_word_file(self):
        """Sign Word document from word_file field by adding user signature"""
        self.ensure_one()
        
        if not DOCX_AVAILABLE:
            raise UserError(_("python-docx library is not installed. Please install it to use document signing."))
        
        # Get user signature
        user = self.env.user
        if not user.sign_signature:
            raise UserError(_("You don't have a signature configured. Please set your signature first."))
        
        # Check if word_file field has data
        if not self.word_file:
            # It's optional, maybe they didn't upload one. We can just return or raise.
            # Usually if they want to sign, they expect a file.
            raise UserError(_("No Word document found to sign."))
        
        # Decode signature image
        try:
            signature_value = user.sign_signature
            # Handle bytes or string format
            if isinstance(signature_value, bytes):
                signature_data = signature_value
            else:
                # Handle data URI format (data:image/png;base64,...)
                if signature_value.startswith('data:'):
                    signature_data = base64.b64decode(signature_value.split(',')[1])
                else:
                    signature_data = base64.b64decode(signature_value)
            
            # Validate signature data
            if not signature_data:
                raise UserError(_("Signature image is empty."))
            
            # Check if it's a valid image (basic check - PNG, JPEG, etc.)
            if len(signature_data) < 10:
                raise UserError(_("Signature image appears to be corrupted or invalid."))
            
            # Detect and convert the image format
            # Check if it's SVG (common in Odoo signatures)
            signature_str = signature_data.decode('utf-8', errors='ignore') if isinstance(signature_data, bytes) else str(signature_data)
            is_svg = signature_str.strip().startswith('<svg') or signature_str.strip().startswith('<?xml')
            
            if is_svg and CAIROSVG_AVAILABLE:
                # Convert SVG to PNG
                try:
                    png_data = cairosvg.svg2png(bytestring=signature_data)
                    signature_data = png_data
                    _logger.info("SVG signature converted to PNG format")
                except Exception as svg_error:
                    _logger.error(f"Error converting SVG signature: {str(svg_error)}", exc_info=True)
                    raise UserError(_("Error converting SVG signature to image. Please update your signature. Error: %s") % str(svg_error))
            elif is_svg and not CAIROSVG_AVAILABLE:
                raise UserError(_("Signature is in SVG format but cairosvg library is not installed. Please install it: pip install cairosvg"))
            
            # Try to validate and convert the image format using PIL if available
            if PIL_AVAILABLE:
                try:
                    # Check if we can identify the image format
                    img_stream = io.BytesIO(signature_data)
                    img_stream.seek(0)
                    
                    # Try to open and identify the image
                    try:
                        img = Image.open(img_stream)
                        img.verify()
                        img_stream.seek(0)
                        img = Image.open(img_stream)
                    except Exception as open_error:
                        # If we can't open it, check if it might be base64 encoded again
                        try:
                            if isinstance(signature_data, bytes):
                                decoded = base64.b64decode(signature_data)
                                if decoded != signature_data:
                                    img_stream = io.BytesIO(decoded)
                                    img_stream.seek(0)
                                    img = Image.open(img_stream)
                                    img.verify()
                                    img_stream.seek(0)
                                    img = Image.open(img_stream)
                                    signature_data = decoded
                                else:
                                    raise open_error
                            else:
                                raise open_error
                        except:
                            raise open_error
                    
                    # Convert to RGB PNG format
                    if img.mode in ('RGBA', 'LA', 'P'):
                        rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'P':
                            img = img.convert('RGBA')
                        if img.mode in ('RGBA', 'LA'):
                            rgb_img.paste(img, mask=img.split()[-1])
                        else:
                            rgb_img.paste(img)
                        img = rgb_img
                    elif img.mode != 'RGB':
                        img = img.convert('RGB')
                    
                    png_stream = io.BytesIO()
                    img.save(png_stream, format='PNG')
                    png_stream.seek(0)
                    signature_data = png_stream.read()
                except Exception as img_error:
                    error_msg = str(img_error)
                    _logger.error(f"Error processing signature image: {error_msg}", exc_info=True)
                    raise UserError(_("Signature image format is not recognized. Please update your signature. Error: %s") % error_msg)
                
        except UserError:
            raise
        except Exception as e:
            raise UserError(_("Error decoding signature: %s") % str(e))
        
        # Process the document
        try:
            doc_data = base64.b64decode(self.word_file)
            # Basic PK check
            if not doc_data.startswith(b'PK'):
                raise UserError(_("The file does not appear to be a valid Word document (.docx)."))
            
            doc_stream = io.BytesIO(doc_data)
            doc = Document(doc_stream)
            
            signature_placeholder = "[ხელმოწერა]"
            signature_found = False
            
            # Helper to search and replace in paragraphs
            def replace_in_paragraph(paragraph):
                nonlocal signature_found
                if signature_placeholder in paragraph.text:
                    signature_found = True
                    text_parts = paragraph.text.split(signature_placeholder)
                    paragraph.clear()
                    if len(text_parts) > 0 and text_parts[0].strip():
                        paragraph.add_run(text_parts[0])
                    
                    signature_stream = io.BytesIO(signature_data)
                    run = paragraph.add_run()
                    try:
                        signature_stream.seek(0)
                        run.add_picture(signature_stream, height=Inches(0.6))
                    except Exception as img_error:
                        raise UserError(_("Error adding signature image: %s") % str(img_error))
                    
                    if len(text_parts) > 1 and text_parts[1].strip():
                        paragraph.add_run(text_parts[1])
                    return True
                return False

            # Check paragraphs
            for paragraph in doc.paragraphs:
                if replace_in_paragraph(paragraph):
                    break
            
            # Check tables if not found
            if not signature_found:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if replace_in_paragraph(paragraph):
                                    break
                            if signature_found: break
                        if signature_found: break
                    if signature_found: break
            
            if not signature_found:
                raise UserError(_("Signature placeholder '[ხელმოწერა]' not found in document."))
            
            output_stream = io.BytesIO()
            doc.save(output_stream)
            output_stream.seek(0)
            
            output_data = output_stream.read()
            self.write({
                'word_file': base64.b64encode(output_data).decode('utf-8'),
            })
            
        except UserError:
            raise
        except Exception as e:
            raise UserError(_("Error processing document: %s") % str(e))
